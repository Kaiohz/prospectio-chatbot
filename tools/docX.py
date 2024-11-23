import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import re




def addText(paragraph : str):
    doc = Document()
    lines = paragraph.split('\n')
    
    for line in lines:
        line = line.strip()
        # Check for headers and format accordingly
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("**") and line.endswith("**"):
            para = doc.add_paragraph()
            run = para.add_run(line[2:-2])
            run.bold = True
        elif line.find("**") != -1:
            para = doc.add_paragraph()
            add_bold_text(para, line)                                                                   
        elif line.startswith("- "):
            doc.add_paragraph(line[2:])
        elif line.startswith("1. "):
            doc.add_paragraph(line[3:])
        elif line:
            doc.add_paragraph(line)
            
    doc.save('demo.docx')
    
def add_bold_text(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        if i % 2 == 0:
            paragraph.add_run(part)
        else:
            paragraph.add_run(part).bold = True
