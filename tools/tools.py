from langchain_core.messages import AIMessage
from langchain_core.tools import tool
from docx import Document
from langgraph.prebuilt import ToolNode
import os
from llm_client.llm_client import LLMProxyChatOpenAI
from memory.memory import messages
import chainlit as cl
import matplotlib.pyplot as plt
from docx.shared import Inches





def load_market_analysis_prompt():
    prompt_path = os.path.join(os.path.dirname(__file__), "../prompts", "market_analysis.txt")
    try:
        with open(prompt_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except FileNotFoundError:
        return "You are a helpful AI assistant."
    
    
llm = LLMProxyChatOpenAI()
market_analysis_prompt = load_market_analysis_prompt()

@tool
def generate_market_report() -> str:
    """Generate market report when its explicitly requested"""
    messages.append({"role": "system", "content": market_analysis_prompt})
    messages.extend(cl.chat_context.to_openai())
    print(messages)
    response =llm.generate([messages])
    filename="market_report.docx"
    paragraph = response.generations[0][0].text
    if os.path.exists(filename):
        doc=Document(filename)
    else:
        doc = Document()
    lines = paragraph.split('\n')
    for line in lines:
        line = line.strip()
        # Check for headers and format accordingly
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("**") and line.endswith("**"):
            para = doc.add_paragraph()
            run = para.add_run(line[2:-2])
            run.bold = True
        elif line.find("**") != -1:
            para = doc.add_paragraph()
            add_bold_text(para, line)                                                                   
        elif line.startswith("- "):
            doc.add_paragraph(line[2:])
        elif line.startswith("1. "):
            doc.add_paragraph(line[3:])
        elif line:
            doc.add_paragraph(line)
    doc.add_picture("./tools/hackathonImage1.png", width=Inches(2), height=Inches(2))
    doc.add_picture("./tools/hackathonImage2.png", width=Inches(2), height=Inches(2))
    doc.add_picture("./tools/hackathonImage3.png", width=Inches(2), height=Inches(2))
    doc.save(filename)
    elements = [
        cl.File(
            name="market.docx",
            path="market_report.docx",
            display="inline",
        ),
    ]

    cl.Message(
        content="Download your market report", elements=elements
    ).send()
    return " "
        

    
def add_bold_text(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        if i % 2 == 0:
            paragraph.add_run(part)
        else:
            paragraph.add_run(part).bold = True


tools = [generate_market_report]
tool_node = ToolNode(tools)